from django.shortcuts import render
from django.http import StreamingHttpResponse, JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
import cv2
import numpy as np
from cvzone.HandTrackingModule import HandDetector
import google.generativeai as genai
from PIL import Image
from docx import Document
from docx.shared import Inches
import os

# Configure the AI model.
genai.configure(api_key="AIzaSyA45efLCmWaiAKizVvJj9QttigH6v9PIQE")
model = genai.GenerativeModel('gemini-1.5-flash')

# Initialize hand detector and webcam.
detector = HandDetector(maxHands=1, modelComplexity=0, detectionCon=0.75, minTrackCon=0.75)
cap = cv2.VideoCapture(0)
if not cap.isOpened():
    raise RuntimeError("Error: Could not open webcam.")

def initialize_canvas(frame):
    return np.zeros_like(frame)

def process_hand(hand):
    lmList = hand["lmList"]
    bbox = hand["bbox"]
    center = hand["center"]
    handType = hand["type"]
    fingers = detector.fingersUp(hand)
    return lmList, bbox, center, handType, fingers

# -----------------------------
# Globals for Gesture Mode
# -----------------------------
gesture_response_text = None
gesture_recognized_equation = None
gesture_solution_text = None

# -----------------------------
# Globals for Speech Mode
# -----------------------------
speech_response_text = None
speech_recognized_equation = None
speech_solution_text = None

# Path for saving an image (used in gesture mode)
equation_image_path = "equation.jpg"

def send_to_ai(model, canvas, fingers):
    global gesture_response_text, gesture_recognized_equation, gesture_solution_text, equation_image_path
    if fingers[4] == 1:
        image = Image.fromarray(canvas)
        image.save(equation_image_path)
        response = model.generate_content(["solve this equation", image])
        if response:
            gesture_response_text = response.text
            gesture_recognized_equation = "Equation from Hand Gesture (see image)"
            gesture_solution_text = gesture_response_text

def save_to_word(equation, solution, image_path=None, filename="equations.docx"):
    try:
        doc = Document()
        doc.add_paragraph(f"Equation: {equation}")
        doc.add_paragraph(f"Solution: {solution}")
        # For gesture mode, include image if available.
        if image_path and os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(4))
        doc.save(filename)
        return filename
    except Exception as e:
        print(f"Error saving to Word: {e}")
        return None

def export_to_word_view(request):
    # Determine which mode's solution to export based on the query parameter "mode".
    mode = request.GET.get("mode", "gesture")
    if mode == "gesture":
        if gesture_recognized_equation and gesture_solution_text:
            filename = save_to_word(gesture_recognized_equation, gesture_solution_text, equation_image_path)
            if filename:
                with open(filename, 'rb') as doc:
                    response = HttpResponse(doc.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    return response
        return JsonResponse({'error': 'No gesture equation found'})
    elif mode == "speech":
        if speech_recognized_equation and speech_solution_text:
            # For speech mode, we pass None as image_path.
            filename = save_to_word(speech_recognized_equation, speech_solution_text, None)
            if filename:
                with open(filename, 'rb') as doc:
                    response = HttpResponse(doc.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    return response
        return JsonResponse({'error': 'No speech equation found'})
    else:
        return JsonResponse({'error': 'Invalid mode'})

def video_stream():
    global gesture_response_text, gesture_recognized_equation, gesture_solution_text, equation_image_path
    prev_pos = None
    drawing = False
    points = []
    smooth_points = None
    _, frame = cap.read()
    canvas = initialize_canvas(frame)

    while True:
        success, img = cap.read()
        if not success:
            break
        img = cv2.flip(img, 1)
        hands, _ = detector.findHands(img, draw=False, flipType=True)
        if hands:
            hand = hands[0]
            lmList, bbox, center, handType, fingers = process_hand(hand)
            index_tip = lmList[8]
            if fingers[1] == 1 and fingers[2] == 0:
                current_pos = np.array([index_tip[0], index_tip[1]])
                if smooth_points is None:
                    smooth_points = current_pos
                else:
                    smooth_points = 0.5 * current_pos + 0.5 * smooth_points
                smoothed_pos = tuple(smooth_points.astype(int))
                if drawing:
                    points.append(smoothed_pos)
                prev_pos = smoothed_pos
                drawing = True
            elif fingers[1] == 1 and fingers[2] == 1:
                drawing = False
                prev_pos = None
                points = []
                smooth_points = None
            elif fingers[0] == 1:
                canvas = initialize_canvas(img)
                points = []
                drawing = False
                prev_pos = None
                smooth_points = None
            elif fingers[4] == 1:
                send_to_ai(model, canvas, fingers)
        if len(points) > 1 and drawing:
            cv2.polylines(canvas, [np.array(points)], isClosed=False, color=(0, 0, 255), thickness=5)
        img = cv2.addWeighted(img, 0.5, canvas, 0.5, 0)
        ret, buffer = cv2.imencode('.jpg', img)
        frame = buffer.tobytes()
        yield (b'--frame\r\n'
               b'Content-Type: image/jpeg\r\n\r\n' + frame + b'\r\n')

@csrf_exempt
def speech_input_view(request):
    global speech_response_text, speech_recognized_equation, speech_solution_text
    if request.method == "POST":
        equation_text = request.POST.get("speech_equation", "")
        if equation_text:
            prompt = f"Solve this math problem: {equation_text}"
            response = model.generate_content([prompt, equation_text])
            if response:
                speech_response_text = response.text
                speech_recognized_equation = equation_text
                speech_solution_text = speech_response_text
            return JsonResponse({'solution': speech_response_text})
        else:
            return JsonResponse({'error': 'No equation received'})
    return JsonResponse({'error': 'Invalid request'})

def index(request):
    return render(request, 'index.html')

def video_feed(request):
    return StreamingHttpResponse(video_stream(), content_type='multipart/x-mixed-replace; boundary=frame')

def get_response(request):
    # This endpoint is used for hand gesture mode to poll the latest solution.
    global gesture_response_text
    return JsonResponse({'response': gesture_response_text})

def process_speech(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            equation = data.get("equation", "")

            # Validate and solve math equation
            solution = solve_math_expression(equation)

            return JsonResponse({"solution": solution})

        except Exception as e:
            return JsonResponse({"error": str(e)})

def solve_math_expression(expression):
    try:
        result = sp.sympify(expression)
        return str(result)  # Ensure clean output
    except Exception:
        return "Invalid Expression"
